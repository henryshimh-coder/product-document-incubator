from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pydantic import BaseModel, ConfigDict
from pypdf import PdfReader

from src.domain.errors import DomainError, ErrorCode
from src.domain.services.file_safety import (
    resolve_source_archive_root,
    validate_business_id,
    validate_upload,
)

MAX_PDF_PAGES = 200
MAX_DOCX_PARAGRAPH_CHARS = 100_000
MAX_CHUNK_CHARS = 2000
CHUNK_OVERLAP_CHARS = 150
_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


class ExtractedChunk(BaseModel):
    model_config = ConfigDict(frozen=True)

    chunk_id: str
    source_id: str
    locator: str
    text: str
    char_start: int
    char_end: int


class ExtractedDocument(BaseModel):
    model_config = ConfigDict(frozen=True)

    source_id: str
    text: str
    chunks: list[ExtractedChunk]


@dataclass(frozen=True)
class _Section:
    locator: str
    text: str


def _read_text(content: bytes) -> str:
    if content.startswith(b"\xef\xbb\xbf"):
        return content.decode("utf-8-sig")
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise DomainError(ErrorCode.EXTRACTION_FAILED, detail="TEXT_DECODE_FAILED") from error


def _extract_pdf(content: bytes) -> list[_Section]:
    reader = PdfReader(BytesIO(content))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise DomainError(ErrorCode.FILE_TOO_LARGE, detail="PDF_PAGE_LIMIT")
    return [
        _Section(locator=f"page:{page_number}", text=page.extract_text() or "")
        for page_number, page in enumerate(reader.pages, start=1)
    ]


def _extract_docx(content: bytes) -> list[_Section]:
    document = Document(BytesIO(content))
    paragraph_chars = sum(len(paragraph.text) for paragraph in document.paragraphs)
    if paragraph_chars > MAX_DOCX_PARAGRAPH_CHARS:
        raise DomainError(ErrorCode.FILE_TOO_LARGE, detail="DOCX_PARAGRAPH_LIMIT")

    title = ""
    sections: list[_Section] = []
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text
        if not text:
            continue
        if paragraph.style.name == "Title" or paragraph.style.name.startswith("Heading"):
            title = text
        locator = f"paragraph:{paragraph_number}"
        if title:
            locator = f"title:{title}; {locator}"
        sections.append(_Section(locator=locator, text=text))
    return sections


def _extract_markdown(text: str) -> list[_Section]:
    headings: dict[int, str] = {}
    sections: list[_Section] = []
    for line_number, line in enumerate(text.splitlines(), start=1):
        match = _MARKDOWN_HEADING.match(line)
        if match:
            level = len(match.group(1))
            headings[level] = match.group(2)
            for deeper_level in tuple(headings):
                if deeper_level > level:
                    del headings[deeper_level]
        if not line:
            continue
        heading_path = " > ".join(headings[level] for level in sorted(headings))
        locator = f"line:{line_number}"
        if heading_path:
            locator = f"heading:{heading_path}; {locator}"
        sections.append(_Section(locator=locator, text=line))
    return sections


def _join_sections(sections: list[_Section]) -> tuple[str, list[tuple[_Section, int]]]:
    positioned: list[tuple[_Section, int]] = []
    cursor = 0
    parts: list[str] = []
    for section in sections:
        if not section.text:
            continue
        if parts:
            cursor += 1
        positioned.append((section, cursor))
        parts.append(section.text)
        cursor += len(section.text)
    return "\n".join(parts), positioned


def _chunk_sections(source_id: str, sections: list[_Section]) -> tuple[str, list[ExtractedChunk]]:
    text, positioned = _join_sections(sections)
    chunks: list[ExtractedChunk] = []
    for section, section_start in positioned:
        offset = 0
        while offset < len(section.text):
            end = min(offset + MAX_CHUNK_CHARS, len(section.text))
            chunks.append(
                ExtractedChunk(
                    chunk_id=f"{source_id}-{len(chunks) + 1:04d}",
                    source_id=source_id,
                    locator=section.locator,
                    text=section.text[offset:end],
                    char_start=section_start + offset,
                    char_end=section_start + end,
                )
            )
            if end == len(section.text):
                break
            offset = end - CHUNK_OVERLAP_CHARS
    return text, chunks


def extract_document_bytes(
    content: bytes,
    *,
    filename: str,
    source_id: str,
) -> ExtractedDocument:
    """Extract text and citation locators from one immutable, caller-supplied byte string."""
    try:
        resolved_source_id = validate_business_id(source_id, "source_id")
        safe_filename = validate_upload(filename, content)
        suffix = Path(safe_filename).suffix.lower()
        if suffix == ".pdf":
            sections = _extract_pdf(content)
        elif suffix == ".docx":
            sections = _extract_docx(content)
        else:
            raw_text = _read_text(content)
            if suffix == ".md":
                sections = _extract_markdown(raw_text)
            else:
                sections = [_Section("text:1", raw_text)]
        text, chunks = _chunk_sections(resolved_source_id, sections)
        if not text.strip() or not chunks:
            raise DomainError(ErrorCode.EXTRACTION_FAILED, detail="EMPTY_EXTRACTED_TEXT")
        return ExtractedDocument(source_id=resolved_source_id, text=text, chunks=chunks)
    except DomainError:
        raise
    except Exception as error:
        raise DomainError(ErrorCode.EXTRACTION_FAILED, detail=type(error).__name__) from error


def extract_document(
    path: Path,
    *,
    source_id: str | None = None,
) -> ExtractedDocument:
    """Read one trusted archive path once, then extract from those immutable bytes."""
    try:
        if source_id is None:
            raise DomainError(ErrorCode.FILE_TYPE_NOT_ALLOWED, detail="SOURCE_ID_REQUIRED")
        resolved_source_id = validate_business_id(source_id, "source_id")
        root = resolve_source_archive_root()
        resolved_path = path.resolve()
        if not resolved_path.is_relative_to(root):
            raise DomainError(ErrorCode.FILE_TYPE_NOT_ALLOWED, detail="UNSAFE_ARCHIVE_PATH")
        relative_path = resolved_path.relative_to(root)
        if len(relative_path.parts) != 3:
            raise DomainError(ErrorCode.FILE_TYPE_NOT_ALLOWED, detail="UNSAFE_ARCHIVE_PATH")
        validate_business_id(relative_path.parts[0], "project_id")
        if relative_path.parts[1] != resolved_source_id:
            raise DomainError(ErrorCode.FILE_TYPE_NOT_ALLOWED, detail="UNSAFE_ARCHIVE_PATH")
        content = resolved_path.read_bytes()
        return extract_document_bytes(
            content,
            filename=resolved_path.name,
            source_id=resolved_source_id,
        )
    except DomainError:
        raise
    except Exception as error:
        raise DomainError(ErrorCode.EXTRACTION_FAILED, detail=type(error).__name__) from error
