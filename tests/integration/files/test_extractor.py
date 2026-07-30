from __future__ import annotations

import importlib
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from src.domain.errors import DomainError, ErrorCode


def extractor_module():
    """Loads the real extraction adapter for an observable missing-feature RED."""
    return importlib.import_module("src.infrastructure.files.extractor")


def _write_pdf(path: Path) -> None:
    stream = b"BT /F1 18 Tf 72 720 Td (PDF extracted content) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, object_body in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode())
        payload.extend(object_body)
        payload.extend(b"\nendobj\n")
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    payload.extend(b"0000000000 65535 f \n")
    payload.extend(b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:]))
    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    )
    payload.extend(trailer.encode())
    path.write_bytes(payload)


@pytest.fixture
def fixture_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.chdir(tmp_path)
    archive_dir = tmp_path / "data" / "source_archive" / "LLD" / "SRC-001"
    archive_dir.mkdir(parents=True)
    _write_pdf(archive_dir / "sample.pdf")
    document = Document()
    document.add_heading("DOCX 标题", level=1)
    document.add_paragraph("DOCX 提取内容")
    document.save(archive_dir / "sample.docx")
    (archive_dir / "sample.txt").write_text("TXT 提取内容", encoding="utf-8")
    (archive_dir / "sample.md").write_text("# 一级标题\n\nMD 提取内容", encoding="utf-8")
    return archive_dir


@pytest.mark.parametrize("fixture_name", ["sample.pdf", "sample.docx", "sample.txt", "sample.md"])
def test_extract_supported_document(fixture_dir: Path, fixture_name: str) -> None:
    """Catches an allowed archive format reaching Ingest without extractable cited text."""
    result = extractor_module().extract_document(
        fixture_dir / fixture_name,
        source_id="SRC-001",
    )

    assert result.text.strip()
    assert result.chunks
    assert all(chunk.locator for chunk in result.chunks)


def test_document_chunks_cap_length_and_overlap(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Catches out-of-bound chunks or a missing overlap losing adjacent source context."""
    monkeypatch.chdir(tmp_path)
    source_text = "a" * 2200
    archive_root = tmp_path / "data" / "source_archive"
    path = archive_root / "LLD" / "SRC-001" / "long.txt"
    path.parent.mkdir(parents=True)
    path.write_text(source_text, encoding="utf-8")

    result = extractor_module().extract_document(
        path,
        source_id="SRC-001",
    )

    assert [(chunk.char_start, chunk.char_end) for chunk in result.chunks] == [
        (0, 2000),
        (1850, 2200),
    ]
    assert [chunk.text for chunk in result.chunks] == [
        source_text[:2000],
        source_text[1850:],
    ]


def test_docx_locator_retains_heading_and_paragraph_number(fixture_dir: Path) -> None:
    """Catches DOCX citations losing the title context needed to locate a paragraph."""
    result = extractor_module().extract_document(
        fixture_dir / "sample.docx",
        source_id="SRC-001",
    )

    assert any("title:DOCX 标题" in chunk.locator for chunk in result.chunks)
    assert any("paragraph:" in chunk.locator for chunk in result.chunks)


def test_extractor_rejects_a_lookalike_source_archive_and_reads_the_fixed_root(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Catches parsing a valid-looking file from a caller-selected source_archive tree."""
    monkeypatch.chdir(tmp_path)
    fixed = tmp_path / "data" / "source_archive" / "LLD" / "SRC-001" / "fixed.txt"
    outside = tmp_path / "untrusted" / "source_archive" / "LLD" / "SRC-001" / "outside.txt"
    fixed.parent.mkdir(parents=True)
    outside.parent.mkdir(parents=True)
    fixed.write_text("受信材料", encoding="utf-8")
    outside.write_text("不可信材料", encoding="utf-8")

    assert extractor_module().extract_document(fixed, source_id="SRC-001").text == "受信材料"
    with pytest.raises(DomainError, match="UNSAFE_ARCHIVE_PATH"):
        extractor_module().extract_document(outside, source_id="SRC-001")


def test_extractor_does_not_accept_a_caller_selected_root(tmp_path: Path) -> None:
    """Catches a public archive_root parameter reintroducing a filesystem trust boundary."""
    path = tmp_path / "untrusted" / "source_archive" / "LLD" / "SRC-001" / "outside.txt"
    path.parent.mkdir(parents=True)
    path.write_text("不可信材料", encoding="utf-8")
    with pytest.raises(TypeError):
        extractor_module().extract_document(
            path,
            source_id="SRC-001",
            archive_root=tmp_path / "untrusted" / "source_archive",
        )


def test_extractor_requires_an_explicit_source_id(fixture_dir: Path) -> None:
    """Catches deriving a business ID from a filename rather than trusted archive metadata."""
    with pytest.raises(DomainError, match="SOURCE_ID_REQUIRED"):
        extractor_module().extract_document(fixture_dir / "sample.txt")


def test_bom_prefixed_text_and_markdown_keep_clean_text_and_heading_locator(
    fixture_dir: Path,
) -> None:
    """Catches a BOM leaking into text or blocking the first Markdown heading locator."""
    (fixture_dir / "bom.txt").write_bytes(b"\xef\xbb\xbf" + "TXT 内容".encode())
    (fixture_dir / "bom.md").write_bytes(b"\xef\xbb\xbf" + "# 一级标题\n正文".encode())

    text_result = extractor_module().extract_document(
        fixture_dir / "bom.txt",
        source_id="SRC-001",
    )
    markdown_result = extractor_module().extract_document(
        fixture_dir / "bom.md",
        source_id="SRC-001",
    )

    assert text_result.text == "TXT 内容"
    assert markdown_result.chunks[0].locator == "heading:一级标题; line:1"


def test_extractor_enforces_pdf_page_limit(fixture_dir: Path) -> None:
    """Catches PDF documents above the 200-page safety limit reaching text extraction."""
    path = fixture_dir / "over-limit.pdf"
    writer = PdfWriter()
    for _ in range(201):
        writer.add_blank_page(width=72, height=72)
    with path.open("wb") as output:
        writer.write(output)

    with pytest.raises(DomainError) as error:
        extractor_module().extract_document(
            path,
            source_id="SRC-001",
        )

    assert error.value.code == ErrorCode.FILE_TOO_LARGE


def test_extractor_enforces_docx_paragraph_character_limit(fixture_dir: Path) -> None:
    """Catches DOCX documents above the 100,000 paragraph-character safety limit."""
    path = fixture_dir / "over-limit.docx"
    document = Document()
    document.add_paragraph("a" * 100_001)
    document.save(path)

    with pytest.raises(DomainError) as error:
        extractor_module().extract_document(
            path,
            source_id="SRC-001",
        )

    assert error.value.code == ErrorCode.FILE_TOO_LARGE


def test_upload_rejects_macro_bearing_docx_and_zip_payload(tmp_path: Path) -> None:
    """Catches macro-bearing OOXML and generic ZIP containers passing upload validation."""
    document = Document()
    document.add_paragraph("safe text")
    docx_path = tmp_path / "safe.docx"
    document.save(docx_path)
    buffer = BytesIO(docx_path.read_bytes())
    with zipfile.ZipFile(buffer, "a") as package:
        package.writestr("word/vbaProject.bin", b"macro")

    with pytest.raises(DomainError, match="MIME_MISMATCH"):
        importlib.import_module("src.domain.services.file_safety").validate_upload(
            "macro.docx", buffer.getvalue()
        )
    with pytest.raises(DomainError, match="UNSAFE_FILENAME"):
        importlib.import_module("src.domain.services.file_safety").validate_upload(
            "payload.zip", b"PK\x03\x04"
        )
